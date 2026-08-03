from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "agreements"
ASSET = ROOT / "videos" / "impactworks-business-automation" / "assets" / "impactworks-logo.png"
OUT.mkdir(parents=True, exist_ok=True)

PRIMARY = "023E8A"
SECONDARY = "0077B6"
TERTIARY = "ABD3FF"
DARK = "071727"
WHITE = "FFFFFF"
INK = "182431"
MUTED = "667384"
RULE = "D7E0EA"
FONT = "SF Pro Text"
FONT_DISPLAY = "SF Pro Display"


def font_file(bold=False, italic=False):
    if bold:
        return "/System/Library/Fonts/SFNS.ttf"
    if italic:
        return "/System/Library/Fonts/SFNSItalic.ttf"
    return "/System/Library/Fonts/SFNS.ttf"


def fit_font(draw, text, max_width, start_size, bold=False):
    size = start_size
    while size > 42:
        f = ImageFont.truetype(font_file(bold=bold), size)
        if draw.textbbox((0, 0), text, font=f)[2] <= max_width:
            return f
        size -= 2
    return ImageFont.truetype(font_file(bold=bold), size)


def white_logo() -> Image.Image:
    src = Image.open(ASSET).convert("RGBA")
    alpha = src.getchannel("A")
    result = Image.new("RGBA", src.size, (255, 255, 255, 0))
    result.putalpha(alpha)
    bbox = result.getbbox()
    return result.crop(bbox) if bbox else result


def cover_image(title: str, subtitle: str, fields: list[tuple[str, str]], slug: str) -> Path:
    out = OUT / f"{slug}-cover.png"
    if out.exists():
        return out
    w, h = 2550, 3300
    img = Image.new("RGB", (w, h), "#" + DARK)
    px = img.load()
    # A restrained deep-blue gradient like the SLA covers.
    for y in range(h):
        for x in range(w):
            glow = max(0.0, 1.0 - (((x - 2250) / 1900) ** 2 + ((y - 520) / 1800) ** 2))
            teal = max(0.0, 1.0 - (((x - 100) / 1500) ** 2 + ((y - 3150) / 1700) ** 2))
            px[x, y] = (
                int(7 + 5 * glow),
                int(23 + 25 * glow + 12 * teal),
                int(39 + 48 * glow + 20 * teal),
            )
    draw = ImageDraw.Draw(img, "RGBA")
    # Approved logomark used as a quiet supporting element.
    mark = Image.open(ASSET).convert("RGBA")
    mark = mark.crop((0, 0, int(mark.width * 0.21), mark.height))
    mark.thumbnail((1450, 1450), Image.Resampling.LANCZOS)
    faint = Image.new("RGBA", mark.size, (255, 255, 255, 0))
    faint.putalpha(mark.getchannel("A").point(lambda a: int(a * 0.055)))
    img.paste(faint, (1450, 1320), faint)

    logo = white_logo()
    logo.thumbnail((1180, 260), Image.Resampling.LANCZOS)
    img.paste(logo, (240, 210), logo)

    kicker = ImageFont.truetype(font_file(bold=True), 52)
    title_font = fit_font(draw, title, 2050, 188, bold=True)
    subtitle_font = fit_font(draw, subtitle, 2050, 78)
    meta_label = ImageFont.truetype(font_file(bold=True), 33)
    meta_value = ImageFont.truetype(font_file(), 48)
    small = ImageFont.truetype(font_file(), 33)

    draw.text((240, 980), "IMPACTWORKS AGREEMENT", font=kicker, fill="#ABD3FF", spacing=8)
    draw.text((240, 1110), title, font=title_font, fill="white")
    draw.text((240, 1415), subtitle, font=subtitle_font, fill="#ABD3FF")
    draw.rounded_rectangle((240, 1590, 570, 1604), radius=7, fill="#0077B6")

    y = 1790
    col_w = 1030
    for i, (label, value) in enumerate(fields):
        col = i % 2
        row = i // 2
        x = 240 + col * col_w
        yy = y + row * 245
        draw.text((x, yy), label.upper(), font=meta_label, fill="#91A8BF")
        draw.text((x, yy + 58), value, font=meta_value, fill="white")

    draw.line((240, 2940, 2310, 2940), fill=(171, 211, 255, 90), width=2)
    draw.text((240, 3000), "Making the internet work for you.", font=small, fill="#ABD3FF")
    draw.text((1985, 3000), "impactworks.com", font=small, fill="#91A8BF")

    img.save(out, dpi=(300, 300), optimize=True)
    return out


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=10.2, color=INK, bold=False, italic=False, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text="", *, bold_prefix=None, after=6, keep=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        style_run(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        style_run(r)
    else:
        r = p.add_run(text)
        style_run(r, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False
    p.add_run(text)
    return p


def add_numbered(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        style_run(p.add_run(item))


def add_bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        style_run(p.add_run(item))


def add_signature_table(doc, left_title, right_title):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.1)
    table.columns[1].width = Inches(3.1)
    cells = table.rows[0].cells
    for cell, title in zip(cells, (left_title, right_title)):
        cell.width = Inches(3.1)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(cell, 100, 110, 100, 110)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(18)
        style_run(p.add_run(title), bold=True, color=PRIMARY)
        for label in ("Signature", "Printed name", "Title", "Date"):
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            style_run(p.add_run(f"{label}: __________________________________"), size=9.6)


def setup_document(cover: Path, short_title: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    p.add_run().add_picture(str(cover), width=Inches(7.5), height=Inches(9.7))
    break_p = doc.add_paragraph()
    break_p.paragraph_format.space_before = Pt(0)
    break_p.paragraph_format.space_after = Pt(0)
    break_p.paragraph_format.line_spacing = Pt(1)

    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    sec2.page_width = Inches(8.5)
    sec2.page_height = Inches(11)
    sec2.top_margin = Inches(0.72)
    sec2.bottom_margin = Inches(0.72)
    sec2.left_margin = Inches(0.82)
    sec2.right_margin = Inches(0.82)
    sec2.header_distance = Inches(0.3)
    sec2.footer_distance = Inches(0.35)
    sec2.different_first_page_header_footer = False

    header = sec2.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    style_run(hp.add_run(f"IMPACTWORKS  /  {short_title.upper()}"), size=7.8, color=MUTED, bold=True)
    footer = sec2.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(fp.add_run("CONFIDENTIAL  |  impactworks.com"), size=7.8, color=MUTED)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, before, after in (("Heading 1", 15, 14, 6), ("Heading 2", 11.5, 10, 4), ("Heading 3", 10.5, 8, 3)):
        st = styles[name]
        st.font.name = FONT_DISPLAY
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_DISPLAY)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_DISPLAY)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(PRIMARY)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = FONT
        st.font.size = Pt(10.2)
        st.paragraph_format.left_indent = Inches(0.38)
        st.paragraph_format.first_line_indent = Inches(-0.18)
    return doc


def opening(doc, title, intro):
    add_heading(doc, title, 1)
    add_paragraph(doc, intro)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    style_run(p.add_run("Agreement reference: "), bold=True, color=PRIMARY)
    style_run(p.add_run("[AGREEMENT REFERENCE]"))
    style_run(p.add_run("    Effective date: "), bold=True, color=PRIMARY)
    style_run(p.add_run("[EFFECTIVE DATE]"))


def build_referral():
    slug = "ImpactWorks-Referral-Partnership-Agreement"
    cover = cover_image(
        "Referral Partnership Agreement",
        "Clear introductions. Shared growth.",
        [("Referral Partner", "[PARTNER LEGAL NAME]"), ("Agreement Ref", "[AGREEMENT REFERENCE]"),
         ("Effective Date", "[EFFECTIVE DATE]"), ("Commission", "10% of collected revenue")], slug)
    doc = setup_document(cover, "Referral Partnership Agreement")
    opening(doc, "Referral Partnership Agreement",
            "This Referral Partnership Agreement (the \"Agreement\") is entered into by Prime Reset LLC, doing business as ImpactWorks, a North Carolina limited liability company with an office at 3600 South College Road, Suite E-125, Wilmington, NC 28412 (\"ImpactWorks\"), and [PARTNER LEGAL NAME], with an address at [PARTNER ADDRESS] (the \"Referral Partner\"). ImpactWorks and the Referral Partner are each a \"Party\" and together the \"Parties.\"")
    add_heading(doc, "1. Purpose and appointment", 1)
    add_paragraph(doc, "ImpactWorks appoints the Referral Partner as a non-exclusive source of introductions for ImpactWorks creative, data, technology, web development, managed services, AI automation, and related offerings (the \"Services\"). The Referral Partner accepts this appointment subject to this Agreement.")
    add_paragraph(doc, "The Referral Partner is an independent contractor. It may not negotiate, promise pricing or performance, sign an agreement, collect funds, or otherwise bind ImpactWorks unless ImpactWorks provides written authority for the specific action.")
    add_heading(doc, "2. Qualified referrals and attribution", 1)
    add_numbered(doc, [
        "A referral must be submitted to ImpactWorks in writing with enough information to identify and contact the prospect.",
        "A prospect is a Qualified Referral only when ImpactWorks confirms acceptance in writing and the prospect was not already an active client, active opportunity, or documented prospect of ImpactWorks during the preceding twelve months.",
        "ImpactWorks may accept or decline any prospect and retains sole authority over proposals, pricing, contracting, delivery, and client acceptance.",
        "If more than one source claims the same prospect, ImpactWorks will determine attribution in good faith using its records and the timing and substance of each introduction."
    ])
    add_heading(doc, "3. Referral fee", 1)
    add_paragraph(doc, "For each Qualified Referral that enters into a paid agreement with ImpactWorks, ImpactWorks will pay the Referral Partner a referral fee equal to ten percent (10%) of amounts actually received by ImpactWorks from that client, excluding taxes, pass-through expenses, refunds, credits, chargebacks, financing charges, and third-party products or licenses (the \"Referral Fee\").")
    add_paragraph(doc, "The Referral Fee applies to the initial agreement and subsequent renewals or extensions entered into while this Agreement remains active, unless the Parties agree otherwise in writing. No fee is earned on unpaid invoices.")
    add_paragraph(doc, "ImpactWorks will calculate earned Referral Fees at the end of each calendar quarter and pay undisputed amounts within thirty (30) days after quarter-end. Each Party is responsible for its own taxes. ImpactWorks may require a completed Form W-9 or other applicable tax documentation before payment.")
    add_heading(doc, "Optional annual incentive", 2)
    add_paragraph(doc, "No annual volume bonus applies unless the Parties complete and sign a written incentive schedule below or in an amendment: [ANNUAL INCENTIVE TERMS, IF ANY].")
    add_heading(doc, "4. Responsibilities and conduct", 1)
    add_bullets(doc, [
        "The Referral Partner will make accurate, lawful introductions and will not misrepresent ImpactWorks or its Services.",
        "The Referral Partner will obtain any consent required before sharing a prospect's personal or business information.",
        "ImpactWorks will contract with, invoice, serve, and support referred clients directly.",
        "Neither Party will offer or accept improper payments, violate anti-bribery rules, send unlawful marketing, or misuse confidential information."
    ])
    add_heading(doc, "5. Confidentiality", 1)
    add_paragraph(doc, "Each Party will protect the other Party's non-public business, technical, financial, client, and commercial information using at least reasonable care; use it only for this relationship; and disclose it only to personnel or advisers who need it and are bound by confidentiality obligations. This section does not apply to information that is publicly available without breach, already known without restriction, independently developed, or lawfully received from another source. Confidentiality obligations continue for five years after disclosure, and for trade secrets as long as applicable law protects them.")
    add_heading(doc, "6. Non-circumvention", 1)
    add_paragraph(doc, "ImpactWorks will not knowingly structure transactions with a Qualified Referral primarily to avoid an earned Referral Fee. This protection continues for two years after the accepted introduction. It does not restrict ImpactWorks from responding to independent inbound requests, serving pre-existing relationships, or working with a prospect after the Referral Partner has been inactive in the opportunity for twelve months.")
    add_heading(doc, "7. Term and termination", 1)
    add_paragraph(doc, "This Agreement begins on the Effective Date and continues until either Party gives thirty (30) days' written notice. A Party may terminate immediately for an uncured material breach after fifteen (15) days' written notice, insolvency, unlawful conduct, misuse of confidential information, or conduct reasonably likely to harm the other Party's reputation. Termination does not eliminate undisputed Referral Fees earned before termination.")
    add_heading(doc, "8. Disclaimers and limitation of liability", 1)
    add_paragraph(doc, "Except for payment obligations, confidentiality breaches, fraud, willful misconduct, or indemnification obligations, neither Party will be liable for indirect, special, incidental, exemplary, or consequential damages. Each Party's aggregate liability under this Agreement will not exceed the Referral Fees paid or payable during the twelve months preceding the event giving rise to the claim.")
    add_heading(doc, "9. General terms", 1)
    add_paragraph(doc, "This Agreement is governed by North Carolina law, without regard to conflicts principles. The state and federal courts located in North Carolina have exclusive jurisdiction. Neither Party may assign this Agreement without the other's written consent, except to a successor in a merger, reorganization, or sale of substantially all relevant assets. Notices must be sent to the addresses or confirmed email addresses stated in this Agreement. This Agreement is the entire agreement on its subject matter and may be amended only in a writing signed by both Parties. If any provision is unenforceable, the remaining provisions remain effective. Electronic signatures and counterparts are valid.")
    doc.add_page_break()
    add_heading(doc, "Signatures", 1)
    add_signature_table(doc, "Prime Reset LLC d/b/a ImpactWorks", "[PARTNER LEGAL NAME]")
    path = OUT / f"{slug}.docx"
    doc.save(path)
    return path


def build_venture():
    slug = "ImpactWorks-Venture-Partnership-Agreement"
    cover = cover_image(
        "Venture Partnership Agreement",
        "Build together. Define it clearly.",
        [("Venture Partner", "[PARTNER LEGAL NAME]"), ("Agreement Ref", "[AGREEMENT REFERENCE]"),
         ("Effective Date", "[EFFECTIVE DATE]"), ("Venture", "[VENTURE NAME / WORKING TITLE]")], slug)
    doc = setup_document(cover, "Venture Partnership Agreement")
    opening(doc, "Venture Partnership Agreement",
            "This Venture Partnership Agreement (the \"Agreement\") is entered into by Prime Reset LLC, doing business as ImpactWorks (\"ImpactWorks\"), and [PARTNER LEGAL NAME], with an address at [PARTNER ADDRESS] (the \"Venture Partner\"). It records the Parties' preliminary relationship for [VENTURE NAME / WORKING TITLE] (the \"Venture\") until replaced by signed formation, operating, shareholder, or similar governing documents.")
    add_heading(doc, "1. Purpose and scope", 1)
    add_paragraph(doc, "The Parties will evaluate, build, launch, and operate the Venture described in [VENTURE DESCRIPTION]. This Agreement applies only to that Venture unless a signed addendum expressly adds another venture.")
    add_heading(doc, "2. Ownership and economics", 1)
    add_paragraph(doc, "The intended ownership split is ImpactWorks: [IMPACTWORKS EQUITY %] and Venture Partner: [PARTNER EQUITY %]. Profit distributions, losses, approved capital requirements, and dilution will follow those percentages unless a later governing document states otherwise.")
    add_paragraph(doc, "Initial contributions are: ImpactWorks - [IMPACTWORKS CONTRIBUTION]; Venture Partner - [PARTNER CONTRIBUTION]. No Party is required to contribute additional funds unless both Parties approve a written capital call stating the amount, purpose, deadline, and effect of non-participation.")
    add_heading(doc, "3. Roles and commitments", 1)
    add_bullets(doc, [
        "ImpactWorks responsibilities: [IMPACTWORKS RESPONSIBILITIES].",
        "Venture Partner responsibilities: [PARTNER RESPONSIBILITIES].",
        "Shared milestones and time commitments: [MILESTONES / TIME COMMITMENTS].",
        "Each Party will maintain accurate records of Venture-related costs, decisions, contracts, credentials, and material work product in mutually accessible systems."
    ])
    add_heading(doc, "4. Decision-making", 1)
    add_paragraph(doc, "Day-to-day decisions may be made by the Party responsible for the relevant function and within the approved budget. The following require both Parties' prior written approval: outside financing; issuance or transfer of ownership; sale or shutdown of the Venture; hiring employees; material contracts; use or licensing of core intellectual property; and any single expenditure above [APPROVAL THRESHOLD] or related expenditures above [MONTHLY THRESHOLD] in a thirty-day period.")
    add_heading(doc, "5. Intellectual property, data, and accounts", 1)
    add_paragraph(doc, "Each Party retains ownership of intellectual property it created before this Agreement or develops independently outside the Venture (\"Background IP\"). To the extent Background IP is required for the Venture, the contributing Party grants the Venture a non-exclusive, non-transferable license during the Term, subject to any written limitations.")
    add_paragraph(doc, "Venture-specific code, designs, documentation, data, AI models, prompts, datasets, fine-tuning assets, brand assets, domains, accounts, and trade secrets created and paid for specifically for the Venture (\"Venture IP\") will be owned by the Venture entity once formed. Before formation, the Parties hold Venture IP for the Venture in proportion to the intended ownership split and may use it only for the Venture. Core repositories, credentials, and documentation must remain accessible to both Parties.")
    add_heading(doc, "6. Confidentiality and restricted conduct", 1)
    add_paragraph(doc, "Each Party will protect non-public Venture and Party information using reasonable care and use it only for the Venture. During the Term and for twelve months after withdrawal, neither Party will knowingly solicit Venture employees, contractors, or active customers to leave the Venture. Any non-competition restriction will apply only if separately stated here and enforceable under applicable law: [OPTIONAL NARROW NON-COMPETE TERMS OR 'NONE'].")
    add_heading(doc, "7. Vesting, withdrawal, and buyout", 1)
    add_paragraph(doc, "Each Party's ownership interest is fully and immediately vested upon execution of this Agreement. A Party may withdraw on thirty (30) days' written notice. The remaining Party or Venture has the first option, exercisable within sixty (60) days, to purchase the departing Party's interest under [BUYOUT FORMULA / VALUATION PROCESS]. If no purchase occurs, the Parties will preserve records, return credentials and property, and cooperate on an orderly transition.")
    add_heading(doc, "8. Formation and definitive documents", 1)
    add_paragraph(doc, "The Parties intend to assign this Agreement, Venture IP, approved contracts, and relevant assets to a properly formed entity. Formation documents should address governance, tax classification, banking, insurance, distributions, deadlock, transfer restrictions, dissolution, and any securities-law requirements. If definitive documents conflict with this Agreement, the definitive documents control.")
    add_heading(doc, "9. Representations, risk, and taxes", 1)
    add_paragraph(doc, "Each Party represents that it has authority to enter into this Agreement and that its contributions will not knowingly infringe third-party rights. Unless otherwise agreed, neither Party may bind the other. Each Party is responsible for its own taxes before formation; the Venture entity will handle entity-level obligations after formation. Neither Party guarantees that the Venture will generate revenue or profit.")
    add_heading(doc, "10. General terms", 1)
    add_paragraph(doc, "This Agreement is governed by North Carolina law, without regard to conflicts principles. The state and federal courts located in North Carolina have exclusive jurisdiction. This Agreement is the entire preliminary agreement concerning the Venture and may be amended only in a signed writing. If any provision is unenforceable, the remainder remains effective. Electronic signatures and counterparts are valid.")
    add_heading(doc, "Signatures", 1)
    add_signature_table(doc, "Prime Reset LLC d/b/a ImpactWorks", "[PARTNER LEGAL NAME]")
    path = OUT / f"{slug}.docx"
    doc.save(path)
    return path


def build_nda():
    slug = "ImpactWorks-Mutual-Non-Disclosure-Agreement"
    cover = cover_image(
        "Mutual Non-Disclosure Agreement",
        "Protect the work before it begins.",
        [("Counterparty", "[COUNTERPARTY LEGAL NAME]"), ("Agreement Ref", "[AGREEMENT REFERENCE]"),
         ("Effective Date", "[EFFECTIVE DATE]"), ("Purpose", "[BUSINESS PURPOSE]")], slug)
    doc = setup_document(cover, "Mutual Non-Disclosure Agreement")
    opening(doc, "Mutual Non-Disclosure Agreement",
            "This Mutual Non-Disclosure Agreement (the \"Agreement\") is entered into by Prime Reset LLC, doing business as ImpactWorks (\"ImpactWorks\"), and [COUNTERPARTY LEGAL NAME], with an address at [COUNTERPARTY ADDRESS] (the \"Counterparty\"). Each may disclose or receive Confidential Information and is a \"Party\"; together they are the \"Parties.\"")
    add_heading(doc, "1. Purpose", 1)
    add_paragraph(doc, "The Parties wish to evaluate or perform [BUSINESS PURPOSE] (the \"Purpose\"). A receiving Party may use the other Party's Confidential Information only for the Purpose.")
    add_heading(doc, "2. Confidential Information", 1)
    add_paragraph(doc, "\"Confidential Information\" means non-public information disclosed in any form that is marked confidential or that a reasonable person would understand to be confidential given its nature and the circumstances. It includes business plans, pricing, proposals, financial information, customer and prospect information, personal data, credentials, systems, source code, designs, documentation, methods, workflows, prompts, models, training data, datasets, recordings, transcripts, research, analyses, and derivative materials.")
    add_heading(doc, "3. Exclusions", 1)
    add_paragraph(doc, "Confidential Information does not include information the receiving Party can document: was publicly available without breach; was already lawfully known without restriction; was lawfully received from another source without a duty of confidentiality; or was independently developed without using the disclosing Party's Confidential Information.")
    add_heading(doc, "4. Protection and permitted disclosure", 1)
    add_numbered(doc, [
        "The receiving Party will use at least reasonable care to protect Confidential Information and no less care than it uses for similar information of its own.",
        "Access may be provided only to employees, contractors, professional advisers, and service providers who need the information for the Purpose and are bound by confidentiality duties at least as protective as this Agreement.",
        "The receiving Party will not disclose, sell, publish, reverse engineer, train an unrelated model on, or otherwise use Confidential Information outside the Purpose without written permission.",
        "If disclosure is legally required, the receiving Party will provide prompt notice when legally permitted, disclose only what is required, and reasonably assist efforts to seek protective treatment."
    ])
    add_heading(doc, "5. Security and incidents", 1)
    add_paragraph(doc, "Each Party will apply reasonable administrative, technical, and physical safeguards appropriate to the sensitivity of the information. A receiving Party will promptly notify the disclosing Party of any known unauthorized access, use, or disclosure and will reasonably cooperate to contain and remediate it.")
    add_heading(doc, "6. Ownership and no license", 1)
    add_paragraph(doc, "Confidential Information remains the property of the disclosing Party. No intellectual-property license or transfer is granted except the limited right to use Confidential Information for the Purpose. Neither Party is obligated to proceed with a transaction, and information is provided without warranty except as stated in a later signed agreement.")
    add_heading(doc, "7. Return or destruction", 1)
    add_paragraph(doc, "Upon written request or completion of the Purpose, the receiving Party will return or securely destroy Confidential Information, except for archival copies maintained automatically or as required by law. Any retained copy remains protected by this Agreement and may not be used for another purpose.")
    add_heading(doc, "8. Term and duration", 1)
    add_paragraph(doc, "This Agreement begins on the Effective Date and covers disclosures made during the following three years. Confidentiality and use restrictions continue for five years after each disclosure. Trade secrets and protected personal data remain protected for as long as applicable law requires or protects them.")
    add_heading(doc, "9. Remedies", 1)
    add_paragraph(doc, "Unauthorized use or disclosure may cause harm that cannot be adequately remedied by money alone. The affected Party may seek injunctive or equitable relief in addition to other available remedies, without waiving any requirement imposed by law.")
    add_heading(doc, "10. General terms", 1)
    add_paragraph(doc, "This Agreement is governed by North Carolina law, without regard to conflicts principles. The state and federal courts located in North Carolina have exclusive jurisdiction. Neither Party may assign this Agreement without the other's written consent, except to a successor in a merger, reorganization, or sale of substantially all relevant assets. This Agreement is the entire agreement concerning confidentiality for the Purpose and may be amended only in a signed writing. If any provision is unenforceable, the remainder remains effective. Electronic signatures and counterparts are valid.")
    add_heading(doc, "Signatures", 1)
    add_signature_table(doc, "Prime Reset LLC d/b/a ImpactWorks", "[COUNTERPARTY LEGAL NAME]")
    path = OUT / f"{slug}.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for result in (build_referral(), build_venture(), build_nda()):
        print(result)
